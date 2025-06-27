#!/usr/bin/env python
# -*- coding: utf-8 -*-

import os
import sys
import json
import logging
from typing import Optional, Dict, Any, List
from dataclasses import dataclass
from datetime import datetime

# MCP imports
from mcp import FastMCP
from mcp.types import Tool, TextContent

# Cross-platform document libraries
try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.enum.style import WD_STYLE_TYPE
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("python-docx not installed. Install with: pip install python-docx", file=sys.stderr)

try:
    import pypandoc
    PANDOC_AVAILABLE = True
except ImportError:
    PANDOC_AVAILABLE = False
    print("pypandoc not installed. Install with: pip install pypandoc", file=sys.stderr)

# Configure logging
log_file = os.path.join(os.path.dirname(os.path.abspath(__file__)), "cross_platform_mcp.log")
logging.basicConfig(
    level=logging.INFO,
    filename=log_file,
    filemode="a",
    format="%(asctime)s - %(name)s - %(levelname)s - %(message)s"
)

# Initialize FastMCP server
mcp = FastMCP("Cross-Platform Document MCP Server")
logger = logging.getLogger(__name__)

# Global document instance
current_document: Optional[Document] = None
document_path: Optional[str] = None

@dataclass
class DocumentInfo:
    """Document metadata"""
    path: str
    created: datetime
    modified: datetime
    format: str

# Helper functions
def get_or_create_document() -> Document:
    """Get current document or create new one"""
    global current_document
    if current_document is None:
        current_document = Document()
    return current_document

def save_current_document(path: str) -> bool:
    """Save current document to file"""
    global document_path
    try:
        doc = get_or_create_document()
        doc.save(path)
        document_path = path
        logger.info(f"Document saved to: {path}")
        return True
    except Exception as e:
        logger.error(f"Error saving document: {e}")
        return False

# MCP Tools
@mcp.tool()
def create_new_document(title: Optional[str] = None) -> str:
    """Create a new document with optional title"""
    global current_document, document_path
    try:
        current_document = Document()
        document_path = None
        
        if title:
            current_document.add_heading(title, level=0)
            
        return "New document created successfully"
    except Exception as e:
        logger.error(f"Error creating document: {e}")
        return f"Error: {str(e)}"

@mcp.tool()
def open_document(file_path: str) -> str:
    """Open an existing document"""
    global current_document, document_path
    try:
        if not os.path.exists(file_path):
            return f"Error: File not found: {file_path}"
            
        current_document = Document(file_path)
        document_path = file_path
        return f"Document opened: {file_path}"
    except Exception as e:
        logger.error(f"Error opening document: {e}")
        return f"Error: {str(e)}"

@mcp.tool()
def save_document(file_path: Optional[str] = None) -> str:
    """Save the current document"""
    global document_path
    try:
        if file_path is None and document_path is None:
            return "Error: No file path specified"
            
        path = file_path or document_path
        if save_current_document(path):
            return f"Document saved to: {path}"
        else:
            return "Error: Failed to save document"
    except Exception as e:
        logger.error(f"Error saving document: {e}")
        return f"Error: {str(e)}"

@mcp.tool()
def insert_text(text: str, style: Optional[str] = None) -> str:
    """Insert text into the document"""
    try:
        doc = get_or_create_document()
        paragraph = doc.add_paragraph(text)
        
        if style:
            paragraph.style = style
            
        return "Text inserted successfully"
    except Exception as e:
        logger.error(f"Error inserting text: {e}")
        return f"Error: {str(e)}"

@mcp.tool()
def insert_heading(text: str, level: int = 1) -> str:
    """Insert a heading (level 1-9)"""
    try:
        doc = get_or_create_document()
        doc.add_heading(text, level=level)
        return f"Heading level {level} inserted"
    except Exception as e:
        logger.error(f"Error inserting heading: {e}")
        return f"Error: {str(e)}"

@mcp.tool()
def insert_table(rows: int, cols: int, data: Optional[List[List[str]]] = None) -> str:
    """Insert a table with optional data"""
    try:
        doc = get_or_create_document()
        table = doc.add_table(rows=rows, cols=cols)
        table.style = 'Table Grid'
        
        if data:
            for i, row_data in enumerate(data[:rows]):
                for j, cell_data in enumerate(row_data[:cols]):
                    table.rows[i].cells[j].text = str(cell_data)
                    
        return f"Table {rows}x{cols} inserted"
    except Exception as e:
        logger.error(f"Error inserting table: {e}")
        return f"Error: {str(e)}"

@mcp.tool()
def insert_image(image_path: str, width_inches: Optional[float] = None) -> str:
    """Insert an image into the document"""
    try:
        if not os.path.exists(image_path):
            return f"Error: Image file not found: {image_path}"
            
        doc = get_or_create_document()
        if width_inches:
            doc.add_picture(image_path, width=Inches(width_inches))
        else:
            doc.add_picture(image_path)
            
        return "Image inserted successfully"
    except Exception as e:
        logger.error(f"Error inserting image: {e}")
        return f"Error: {str(e)}"

@mcp.tool()
def export_to_pdf(output_path: str) -> str:
    """Export document to PDF (requires pypandoc)"""
    if not PANDOC_AVAILABLE:
        return "Error: pypandoc not installed. Install with: pip install pypandoc"
        
    try:
        if document_path is None:
            temp_path = "temp_document.docx"
            save_current_document(temp_path)
        else:
            temp_path = document_path
            
        pypandoc.convert_file(temp_path, 'pdf', outputfile=output_path)
        
        if document_path is None and os.path.exists(temp_path):
            os.remove(temp_path)
            
        return f"Document exported to PDF: {output_path}"
    except Exception as e:
        logger.error(f"Error exporting to PDF: {e}")
        return f"Error: {str(e)}"

@mcp.tool()
def find_and_replace(find_text: str, replace_text: str) -> str:
    """Find and replace text in the document"""
    try:
        doc = get_or_create_document()
        count = 0
        
        for paragraph in doc.paragraphs:
            if find_text in paragraph.text:
                paragraph.text = paragraph.text.replace(find_text, replace_text)
                count += 1
                
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if find_text in cell.text:
                        cell.text = cell.text.replace(find_text, replace_text)
                        count += 1
                        
        return f"Replaced {count} occurrences"
    except Exception as e:
        logger.error(f"Error in find and replace: {e}")
        return f"Error: {str(e)}"

@mcp.tool()
def set_font(font_name: str, font_size: Optional[int] = None) -> str:
    """Set default font for new paragraphs"""
    try:
        doc = get_or_create_document()
        style = doc.styles['Normal']
        font = style.font
        font.name = font_name
        
        if font_size:
            font.size = Pt(font_size)
            
        return f"Font set to {font_name}" + (f" {font_size}pt" if font_size else "")
    except Exception as e:
        logger.error(f"Error setting font: {e}")
        return f"Error: {str(e)}"

@mcp.tool()
def add_page_break() -> str:
    """Add a page break"""
    try:
        doc = get_or_create_document()
        doc.add_page_break()
        return "Page break added"
    except Exception as e:
        logger.error(f"Error adding page break: {e}")
        return f"Error: {str(e)}"

@mcp.tool()
def get_document_info() -> str:
    """Get information about the current document"""
    try:
        if current_document is None:
            return "No document is currently open"
            
        info = {
            "path": document_path or "Unsaved",
            "paragraphs": len(current_document.paragraphs),
            "tables": len(current_document.tables),
            "sections": len(current_document.sections)
        }
        
        return json.dumps(info, indent=2)
    except Exception as e:
        logger.error(f"Error getting document info: {e}")
        return f"Error: {str(e)}"

@mcp.tool()
def list_available_styles() -> str:
    """List all available paragraph styles"""
    try:
        doc = get_or_create_document()
        styles = [style.name for style in doc.styles if style.type == WD_STYLE_TYPE.PARAGRAPH]
        return json.dumps(styles, indent=2)
    except Exception as e:
        logger.error(f"Error listing styles: {e}")
        return f"Error: {str(e)}"

@mcp.tool()
def insert_bullet_list(items: List[str]) -> str:
    """Insert a bulleted list"""
    try:
        doc = get_or_create_document()
        for item in items:
            p = doc.add_paragraph(item, style='List Bullet')
        return f"Bullet list with {len(items)} items inserted"
    except Exception as e:
        logger.error(f"Error inserting bullet list: {e}")
        return f"Error: {str(e)}"

@mcp.tool()
def insert_numbered_list(items: List[str]) -> str:
    """Insert a numbered list"""
    try:
        doc = get_or_create_document()
        for item in items:
            p = doc.add_paragraph(item, style='List Number')
        return f"Numbered list with {len(items)} items inserted"
    except Exception as e:
        logger.error(f"Error inserting numbered list: {e}")
        return f"Error: {str(e)}"

def main():
    """Main entry point"""
    try:
        logger.info("Starting Cross-Platform Document MCP Server")
        
        if not DOCX_AVAILABLE:
            logger.error("python-docx is required. Install with: pip install python-docx")
            sys.exit(1)
            
        mcp.run()
    except Exception as e:
        logger.error(f"Server error: {e}")
        sys.exit(1)

if __name__ == "__main__":
    main()